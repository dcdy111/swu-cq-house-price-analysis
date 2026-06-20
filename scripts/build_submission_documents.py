from __future__ import annotations

import argparse
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
TEMPLATE = WORKSPACE / "学年设计Ⅱ文档模板.docx"
EVIDENCE = ROOT / "docs" / "evidence" / "acceptance_evidence.json"
EVIDENCE_DIR = ROOT / "docs" / "evidence"
OUTPUT_DIR = ROOT / "deliverables"
TITLE = "重庆市二手房源价格数据分析与智能可视化系统"
STUDENT = "张浩博"
ACCENT = "173F73"
ACCENT_LIGHT = "EAF0F8"
GRAY = "F3F4F6"


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 10.5, bold=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc: Document, compact: bool = False) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5 if not compact else 10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5 if not compact else 1.25
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0 if not compact else 4)

    for name, size, before, after, level in (
        ("Heading 1", 16, 16, 10, 1),
        ("Heading 2", 14, 12, 6, 2),
        ("Heading 3", 12, 8, 4, 3),
    ):
        try:
            style = doc.styles[name]
        except KeyError:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style_pr = style.element.get_or_add_pPr()
        outline = style_pr.find(qn("w:outlineLvl"))
        if outline is None:
            outline = OxmlElement("w:outlineLvl")
            style_pr.append(outline)
        outline.set(qn("w:val"), str(level - 1))
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string("000000" if not compact else ACCENT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 and not compact else WD_ALIGN_PARAGRAPH.LEFT

    for name in ("List Bullet", "List Number"):
        if name in doc.styles:
            style = doc.styles[name]
            style.font.name = "Times New Roman"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            style.font.size = Pt(10.5 if not compact else 10)
            style.paragraph_format.line_spacing = 1.35 if not compact else 1.2
            style.paragraph_format.space_after = Pt(2)


def set_page(section) -> None:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_header_footer(doc: Document, label: str) -> None:
    for index, section in enumerate(doc.sections):
        set_page(section)
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if index > 0:
            r = p.add_run(label)
            set_run_font(r, east_asia="宋体", size=9, color="666666")
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if index > 0:
            r = fp.add_run("第 ")
            set_run_font(r, east_asia="宋体", size=9, color="666666")
        if index > 0:
            pg_num_type = section._sectPr.find(qn("w:pgNumType"))
            if pg_num_type is None:
                pg_num_type = OxmlElement("w:pgNumType")
                section._sectPr.append(pg_num_type)
            pg_num_type.set(qn("w:start"), "1")
            add_page_field(fp)
            r = fp.add_run(" 页")
            set_run_font(r, east_asia="宋体", size=9, color="666666")


def clear_cell(cell) -> None:
    cell.text = ""
    if not cell.paragraphs:
        cell.add_paragraph()


def fill_cover(doc: Document) -> None:
    if len(doc.paragraphs) > 4:
        p = doc.paragraphs[4]
        p.text = "学年设计Ⅱ文档"
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            set_run_font(p.runs[0], east_asia="黑体", size=22, bold=True)
    table = doc.tables[0]
    values = [
        TITLE,
        STUDENT,
        "",
        "",
        "",
        "",
        "2026年6月",
    ]
    for row, value in zip(table.rows, values):
        clear_cell(row.cells[1])
        p = row.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(value)
        set_run_font(r, east_asia="宋体", size=11, bold=(value == TITLE))


def trim_template_body(doc: Document) -> None:
    start = next((p for p in doc.paragraphs if p.text.strip().startswith("第1章 概述")), None)
    if start is None:
        raise RuntimeError("模板中未找到正文起始位置")
    toc = next((p for p in doc.paragraphs if p.text.strip() == "请生成目录"), None)
    if toc is None:
        raise RuntimeError("模板中未找到目录占位符")
    toc.text = "[[TOC]]"
    toc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body = doc.element.body
    seen_toc = False
    for child in list(body):
        if child is toc._p:
            seen_toc = True
            continue
        if seen_toc and child.tag != qn("w:sectPr"):
            body.remove(child)
    doc.add_section(WD_SECTION.NEW_PAGE)


def add_heading(doc: Document, text: str, level: int = 1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str, indent: bool = True):
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21) if indent else Pt(0)
    p.paragraph_format.keep_together = False
    r = p.add_run(text.strip())
    set_run_font(r)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.first_line_indent = None
    r = p.add_run(text)
    set_run_font(r)
    return p


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list], widths_dxa: list[int], caption: str | None = None):
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(caption)
        set_run_font(r, east_asia="黑体", size=10.5, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths_dxa)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, ACCENT_LIGHT)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(header))
        set_run_font(r, east_asia="黑体", size=9.5, bold=True, color=ACCENT)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 or len(str(value)) < 16 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc: Document, image_path: Path, caption: str, width_in: float = 6.0):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_in))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(3)
    cp.paragraph_format.space_after = Pt(8)
    cp.paragraph_format.keep_with_next = False
    r = cp.add_run(caption)
    set_run_font(r, east_asia="宋体", size=9, color="555555")


def add_code(doc: Document, code: str, caption: str | None = None):
    if caption:
        p = doc.add_paragraph()
        r = p.add_run(caption)
        set_run_font(r, east_asia="黑体", size=10, bold=True)
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_geometry(table, [9120])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F8FA")
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.05
    for idx, line in enumerate(code.strip().splitlines()):
        if idx:
            p.add_run("\n")
        r = p.add_run(line)
        set_run_font(r, east_asia="等线", latin="Consolas", size=8.5)


def load_font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf" if bold else "C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def create_evidence_cards(evidence: dict) -> None:
    EVIDENCE_DIR.mkdir(parents=True, exist_ok=True)
    cards = [
        (
            EVIDENCE_DIR / "05-database-evidence.png",
            "MySQL 数据底座验收快照",
            [
                "SELECT COUNT(*) FROM listings WHERE status IN ('active','valid');",
                f"→ {evidence['database']['valid_listings']:,} 条有效房源，覆盖 {evidence['database']['district_count']} 个区县",
                "SELECT source, COUNT(*) FROM listings GROUP BY source;",
                "→ " + "，".join(f"{k}={v:,}" for k, v in evidence["database"]["source_counts"].items()),
                "SELECT COUNT(*) FROM listing_snapshots;",
                f"→ {evidence['database']['snapshots']:,} 条快照；{evidence['database']['listings_with_multiple_snapshots']} 套房源存在多次快照",
                f"质量报告 #{evidence['quality']['id']}：可分析 {evidence['quality']['analysis_ready_count']:,} 条，极值 {evidence['quality']['extreme_count']} 条",
            ],
        ),
        (
            EVIDENCE_DIR / "06-agent-evidence.png",
            "DeepSeek Agent 工具调用验收快照",
            [
                f"agent_tool_calls：{evidence['agent']['tool_call_count']} 条",
                f"generated_reports：{evidence['agent']['generated_report_count']} 份",
                f"最新报告：#{evidence['agent']['latest_report']['id']} {evidence['agent']['latest_report']['title']}",
                "证据键：" + "、".join(evidence['agent']['latest_report']['evidence_keys']),
                "最近工具：" + " → ".join(row['tool_name'] for row in evidence['agent']['recent_calls'][:5]),
                "报告 PDF：output/pdf/重庆二手房挂牌价分析_Agent报告证据_报告5.pdf",
                "口径约束：所有价格均表述为挂牌价/报价，Agent 不执行用户 SQL。",
            ],
        ),
    ]
    for path, title, lines in cards:
        image = Image.new("RGB", (1600, 820), "#0C1220")
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle((55, 50, 1545, 770), radius=28, fill="#111827", outline="#334155", width=3)
        draw.text((100, 92), title, font=load_font(42, True), fill="#E8F0FF")
        draw.text((100, 155), f"生成时间：{evidence['generated_at']}｜来源：acceptance_evidence.json", font=load_font(23), fill="#94A3B8")
        y = 225
        for line in lines:
            draw.text((112, y), line, font=load_font(27), fill="#D8E2F2")
            y += 70
        image.save(path)


def p(doc: Document, *paragraphs: str):
    for text in paragraphs:
        add_body(doc, text)


def build_main(e: dict) -> Path:
    doc = Document(str(TEMPLATE))
    fill_cover(doc)
    trim_template_body(doc)
    configure_styles(doc, compact=False)
    configure_header_footer(doc, TITLE + "｜学年设计说明书")

    db = e["database"]
    quality = e["quality"]
    job = e["analysis_job"]
    regression = e["model_results"]["regression"]
    cluster = e["model_results"]["cluster"]
    anomaly = e["model_results"]["anomaly"]

    add_heading(doc, "摘  要", 1)
    p(
        doc,
        f"本项目围绕重庆市二手房挂牌价数据的获取、治理、分析与展示，完成了一套轻量化后台管理式数据产品。系统以 MySQL 8.x 为唯一运行数据库，后端采用 Flask 分层 API，前端采用 React、TypeScript 与 Vite，图表由 Recharts 和 ECharts 共同完成，数据分析使用 scikit-learn，智能问数与报告生成接入 DeepSeek。与一次性爬取后静态画图的方案相比，本项目更重视数据来源说明、增量快照、质量监控、模型证据和 Agent 工具调用记录，形成“数据源—采集任务—清洗入库—增量维护—分析建模—可视化—智能报告”的闭环。",
        f"截至 2026 年 6 月 18 日，业务库共有 {db['valid_listings']:,} 条有效房源，覆盖 {db['district_count']} 个区县，其中旧库冷启动数据 {db['source_counts']['anjuke_legacy'] + db['source_counts']['lianjia_legacy']:,} 条，新系统房天下标准采集数据 {db['source_counts']['fang']} 条；快照表共有 {db['snapshots']:,} 条记录。最新分析任务采用按区县分层、固定顺序抽样的 5,000 条样本，覆盖全部 38 个有效区县和 3 类来源。随机森林回归的 MAE 为 {regression['metrics']['mae']} 元/平方米、RMSE 为 {regression['metrics']['rmse']} 元/平方米、R² 为 {regression['metrics']['r2']}；KMeans 轮廓系数为 {cluster['metrics']['silhouette_score']}；IsolationForest 与强规则联合识别 {anomaly['metrics']['anomaly_count']} 条异常样本，异常率为 {anomaly['metrics']['anomaly_rate'] * 100:.2f}%。",
        f"系统已通过 Python 语法检查、34 项后端自动化测试和前端生产构建，并使用真实业务库完成 Dashboard、采集任务、分析建模等页面验收。DeepSeek Agent 已保存 {e['agent']['tool_call_count']} 条工具调用和 {e['agent']['generated_report_count']} 份报告。本文只将上述结果表述为对挂牌价的描述、解释和辅助估计，不将其等同于成交价，也不把冷启动旧数据描述为新系统实时采集数据。部署工作在本轮暂缓，文中仅给出后续部署方案。",
    )
    add_body(doc, "关键词：重庆二手房；挂牌价；增量快照；数据质量；机器学习；智能问数", indent=False)

    add_heading(doc, "第1章 概述", 1)
    add_heading(doc, "1.1 项目背景", 2)
    p(
        doc,
        "重庆市辖区范围广，不同区县在产业结构、交通条件、公共服务、建设年代和住房供给上存在明显差异。公开房源平台能够持续产生二手房挂牌信息，但不同平台字段命名、单位、页面结构和反爬策略不一致，同一房源还会出现重复展示、价格变化、状态变化和短期下架等情况。如果只把一次采集结果保存为 CSV，再制作若干静态图表，数据很快失去时效，也无法说明模型结论来自哪一批样本。",
        "学院指导手册要求数据不少于 5 万条，完成清洗整理与存储，支持定期增量更新，以 Web 或 App 展示可视化，并使用挖掘算法得到相关结论。项目因此把重点从“爬多少数据”扩展为“如何形成可信、可维护、可展示的数据系统”。最终实现不仅包含采集脚本和图表，还包含任务状态、失败日志、唯一指纹、价格快照、质量报告、分析任务、模型结果和 Agent 工具证据。",
    )
    add_heading(doc, "1.2 研究对象与口径", 2)
    p(
        doc,
        "本项目研究对象是公开页面中的重庆市二手房挂牌价或报价。挂牌价反映业主和中介在特定时间点的市场预期，能够用于区域比较、结构分析和异常识别，但不能直接等同于真实成交价。系统前端、接口、报告和本文统一使用“挂牌价”“挂牌单价”“挂牌总价”等表述，避免因概念混用而夸大结论。",
        "数据由两层构成：第一层是旧项目 SQLite 数据导入形成的冷启动基线，用于达到课程要求的数据规模并提供区县覆盖；第二层是当前系统通过房天下等来源采集的标准化增量数据，用于证明采集、去重、last_seen_at 更新、价格快照和失败日志链路。两层数据都进入 MySQL，但通过 source 字段保留来源，分析结果同时记录来源分布，防止把 12 万条历史样本描述成新系统实时抓取。",
    )
    add_heading(doc, "1.3 项目目标", 2)
    p(
        doc,
        "数据目标是构建字段统一、来源可追踪、数量超过 5 万条的重庆二手房挂牌价数据底座；工程目标是实现可查询、可采集、可分析、可问数的前后端系统；研究目标是使用描述统计、回归、聚类和异常检测回答区域差异、价格驱动因素、市场分层和异常挂牌等问题；展示目标是让答辩教师能够通过页面、数据库记录、任务日志和测试命令核验系统功能。",
        "项目坚持最小但完整的交付原则。第一版不引入微服务、复杂权限中心或大数据集群，而是使用 Flask、MySQL 和 React 完成全链路；部署不作为本轮完成条件，避免在核心数据与文档尚未收口时分散精力。功能完成后再补充一键启动脚本、增量验收脚本、验收证据导出脚本和正式文档生成脚本，降低后续复现成本。",
    )
    add_heading(doc, "1.4 主要成果与创新点", 2)
    p(
        doc,
        "第一，系统把采集过程建模为可持久化任务，记录 pending、running、success、partial_failed 等状态以及成功页、失败页、解析条数和错误类型，使反爬波动不再只是终端中的临时错误。第二，唯一指纹不包含价格，同一房源的价格变化会更新主表并写入 snapshot，而不是重复新增房源。第三，分析任务由“最近 N 条”改为按区县分层的确定性抽样，显著降低只覆盖少数区县造成的选择偏差。",
        "第四，异常检测不直接把 IsolationForest 的所有孤立样本都当成错误，而是结合区县中位数偏离、Z 分数和强规则给出异常理由，并把异常记录保留供复核。第五，Agent 不直接连接 MySQL，也不执行用户输入 SQL，而是调用 ToolRegistry 中的白名单工具；具体数值来自工具 JSON，工具参数、结果和耗时保存到 agent_tool_calls，报告证据保存到 generated_reports.evidence_json。",
    )

    add_heading(doc, "第2章 系统分析和可行性分析", 1)
    add_heading(doc, "2.1 学院要求映射", 2)
    add_table(
        doc,
        ["学院要求", "系统实现", "真实证据", "状态"],
        [
            ["不少于5万条", "MySQL listings 统一存储", f"有效 {db['valid_listings']:,} 条，38区县", "完成"],
            ["数据清洗整理", "字段解析、标准化、质量分与异常标记", f"可分析 {quality['analysis_ready_count']:,} 条", "完成"],
            ["定期增量", "APScheduler 注册、手动增量接口、指纹 upsert", "重复采集153条未变；真实价格变化1条", "代码完成，常驻暂缓"],
            ["Web可视化", "Dashboard、房源、任务、质量、模型、Agent", "真实后端页面截图", "完成"],
            ["挖掘结论", "回归、KMeans、IsolationForest", f"R²={regression['metrics']['r2']}，异常率={anomaly['metrics']['anomaly_rate']*100:.2f}%", "完成"],
            ["稳定美观", "统一异常响应、任务日志、科研简约界面", "34项测试、构建和浏览器检查", "完成"],
        ],
        [1800, 2700, 2700, 1920],
        "表2-1 学院要求与系统实现映射",
    )
    p(
        doc,
        "映射结果表明，部署暂缓并不影响学院对数据、清洗、增量、可视化和挖掘的核心要求。需要明确的是，“定期更新”在非部署环境中无法通过服务器连续运行数天来证明，因此本轮采用三类证据：调度器注册逻辑和接口测试、手动触发增量任务、重复采集与价格变化的数据库结果。后续上线后再开启 SCHEDULER_ENABLED 和增量任务开关。",
        "说明书、安装使用说明、设计日志和答辩证据清单分别生成，避免将所有内容堆在一份文档中。主说明书保留项目背景、分析设计、实现过程、关键代码、问题解决和总结展望；安装文档给出环境变量、数据库初始化和启动流程；日志按日期记录可核验活动；证据清单逐项标明截图或命令位置。",
    )
    add_heading(doc, "2.2 功能需求", 2)
    p(
        doc,
        "系统用户以课程答辩和研究演示为主，核心操作包括登录后台、查看总体指标和区县图表、按条件检索房源、创建小规模采集任务、查看失败日志、生成质量报告、启动分析任务、查看模型指标和异常房源、向 Agent 提问并导出报告。所有接口返回统一 code、message、data 和 trace_id，便于前端提示与错误定位。",
        "房源查询支持 district、price_min、price_max、area_min、area_max、keyword、page 和 page_size。采集任务支持来源、区县、页数、并发和模式参数。分析任务支持 EDA、回归、聚类、异常检测或 all。Agent 支持市场统计、图表序列、采集状态、增量任务、分析任务、模型结果和报告生成等白名单工具，禁止任意 SQL。",
    )
    add_heading(doc, "2.3 非功能需求", 2)
    p(
        doc,
        "稳定性方面，爬虫单页失败不会中断整个任务，失败页进入日志并允许任务以 partial_failed 收口；后端异常通过统一响应返回，前端展示加载态和错误提示。性能方面，常用图表使用数据库聚合与有限样本查询，在 12 万条数据下本机实测均低于 2 秒，其中 overview 约 1.02 秒，区县均价约 0.25 秒，价格分布约 0.31 秒。",
        "安全与伦理方面，密钥保存在环境变量或数据库秘密配置中，不写入仓库；Agent 只通过 ToolRegistry 访问服务层；爬虫默认低并发、随机间隔，不绕过验证码，不做强对抗反爬；公开数据仅用于课程研究。可维护性方面，后端保持 api—services—models—crawlers/tasks/ml/agent 分层，路由不承载复杂 SQL 和模型训练。",
    )
    add_heading(doc, "2.4 可行性与风险", 2)
    p(
        doc,
        "技术可行性来自成熟的 Python 与 Web 生态。Flask 适合课程项目快速构建 API；SQLAlchemy 与 PyMySQL 能对接 MySQL 8.x；ThreadPoolExecutor 足以完成低并发分片采集；Pandas 与 scikit-learn 可完成特征工程和模型训练；React 与图表库能提供清晰交互。项目无需 GPU 和分布式集群，普通开发电脑即可运行。",
        "主要风险包括页面结构变化、反爬网关、旧数据与新数据口径差异、模型抽样偏差、趋势时间跨度不足和大模型幻觉。对应措施是保留失败 URL 与错误类型、将旧数据标注为 legacy 来源、记录抽样策略与来源分布、只基于真实 snapshot_time 绘制趋势、用工具白名单约束 Agent，并在结论中展示误差指标和数据限制。",
    )

    add_heading(doc, "第3章 系统设计", 1)
    add_heading(doc, "3.1 总体架构", 2)
    p(
        doc,
        "系统采用前后端分离但部署简单的单体架构。浏览器访问 React 页面，前端通过统一 API 客户端请求 Flask；Blueprint 只负责参数接收和统一响应，Service 处理查询、采集、质量、分析和 Agent 业务，SQLAlchemy 模型映射 MySQL 表。采集任务由爬虫注册表选择具体来源，分析服务调用 scikit-learn，Agent 服务调用白名单工具后再请求 DeepSeek。",
        "数据流从公开房源页面开始，经解析与字段标准化后进入 ListingService.upsert_listing。新房源写主表和初始快照，未变房源只更新 last_seen_at，价格变化更新当前价格并追加快照。Dashboard 从服务层聚合 listings 与 listing_snapshots；分析任务把抽样数据转换为特征矩阵并保存 metrics、artifacts 和 evidence；Agent 工具读取这些聚合结果，不直接暴露数据库。",
    )
    add_table(
        doc,
        ["层级", "主要技术", "职责"],
        [
            ["前端", "React + TypeScript + Vite", "页面状态、交互、API调用"],
            ["图表", "Recharts + ECharts", "常规统计图与重庆地图"],
            ["后端", "Flask + Blueprint + Service", "接口、业务编排、统一响应"],
            ["数据", "MySQL 8.x + SQLAlchemy + PyMySQL", "唯一运行数据库与持久化"],
            ["采集", "requests + BeautifulSoup + ThreadPoolExecutor", "低并发、多来源、失败日志"],
            ["分析与智能", "Pandas + scikit-learn；DeepSeek + ToolRegistry", "EDA、回归、聚类、异常检测、问数与报告"],
        ],
        [1600, 2700, 4820],
        "表3-1 最终技术栈",
    )
    add_heading(doc, "3.2 数据库设计", 2)
    p(
        doc,
        "数据库以 listings 为中心，周边表分别记录快照、采集、质量、模型和 Agent 证据。listings 保存当前状态，listing_snapshots 保存时间序列；crawl_tasks 与 crawl_logs 构成任务审计；data_quality_reports 保存一次质量扫描；analysis_jobs 与 model_results 保存训练任务和结果；agent_tool_calls 与 generated_reports 保存智能问答证据。",
        "主表使用 source 与 fingerprint 联合唯一约束。fingerprint 由来源房源标识、规范化链接、标题、区县、小区、面积和户型等稳定字段构成，不包含价格。这样价格变化不会破坏同一性。连续 N 次未出现标记 inactive 的策略在当前有限页补采条件下暂不启用，避免把未覆盖页面中的房源误判为下架。",
    )
    add_table(
        doc,
        ["表名", "核心作用", "关键证据"],
        [
            ["listings", "标准房源当前状态", f"{db['listings_total']:,} 条"],
            ["listing_snapshots", "价格与状态时间快照", f"{db['snapshots']:,} 条"],
            ["crawl_tasks / crawl_logs", "任务状态与页级日志", "成功、部分失败、拦截原因"],
            ["data_quality_reports", "质量扫描快照", f"最新报告 #{quality['id']}"],
            ["analysis_jobs / model_results", "模型任务与证据", f"最新任务 #{job['id']}"],
            ["agent_tool_calls", "工具参数、结果和耗时", f"{e['agent']['tool_call_count']} 条"],
            ["generated_reports", "报告正文与 evidence_json", f"{e['agent']['generated_report_count']} 份"],
        ],
        [2400, 3650, 3070],
        "表3-2 核心数据表",
    )
    add_figure(doc, EVIDENCE_DIR / "05-database-evidence.png", "图3-1 MySQL 数据底座验收脚本输出快照")
    add_heading(doc, "3.3 API 与模块边界", 2)
    p(
        doc,
        "后端入口 create_app 负责加载配置、初始化扩展和注册 Blueprint。api 目录只做参数转换和返回；services 目录集中业务逻辑；models 目录定义 ORM；crawlers 目录封装数据源；tasks 目录负责调度；agent 目录封装 ToolRegistry 与 DeepSeek 客户端。该结构使房源查询、采集和模型训练可以分别测试，也避免在路由中拼接复杂 SQL。",
        "接口统一返回 JSON 结构并生成 trace_id。健康检查、概览、图表、房源查询、质量、采集、调度、分析、设置和 Agent 均通过 Blueprint 注册。前端 API 客户端负责附加认证 Token、解析错误与类型转换。登录采用课程演示所需的本地管理员 Token，不声称已经接入学校统一身份认证。",
    )
    add_code(
        doc,
        """{
  "code": 0,
  "message": "ok",
  "data": {},
  "trace_id": "20260618-abcdef12"
}""",
        "统一接口返回示例",
    )
    add_heading(doc, "3.4 前端信息架构", 2)
    p(
        doc,
        "前端保留现有 React + TypeScript 方案，不为追求与早期需求文档一致而重写为原生 JavaScript。侧边栏包含首页总览、房源数据管理、采集任务管理、数据清洗质量、分析建模、智能问答与报告、系统设置七个模块。页面使用白底、深蓝标题、浅灰分区与少量橙色强调，优先展示真实数据状态而非复杂动效。",
        "Dashboard 负责总体 KPI、区县地图与排行、总价分布、面积—单价散点、快照趋势、户型分布和任务状态。分析页展示 MAE、RMSE、R²、MAPE、特征重要性、箱线图、聚类、异常和模型对比，并新增区县覆盖、来源数和异常率。登录页技术口径已改为 scikit-learn、DeepSeek、ECharts/Recharts，同时移除未实现的 SSO 和忘记密码入口。",
    )
    add_figure(doc, EVIDENCE_DIR / "01-dashboard.png", "图3-2 Dashboard 真实业务库页面")

    add_heading(doc, "第4章 数据获取与实现", 1)
    add_heading(doc, "4.1 数据源与采集边界", 2)
    p(
        doc,
        "系统实现房天下、安居客移动端和链家等爬虫适配，其中房天下当前可直接返回普通 HTML，作为默认小规模采集源；安居客移动端页面结构存在波动；链家需要 Cookie，仅作为实验来源。采集器遵循确定性程序控制，Agent 不负责打开网页或绕过验证码。默认并发为 3—5，请求间隔 1—3 秒，页数和区县均由任务参数限制。",
        "BaseCrawler 统一请求、限速和错误包装，具体爬虫只实现 URL 构造与页面解析。CrawlerRegistry 根据 source 选择实现，CrawlService 创建任务、运行分片、写日志并调用 ListingService 入库。页面被登录、验证码或反爬网关拦截时，记录明确错误而不是把空页面当成成功。失败页不会导致 Flask 服务崩溃。",
    )
    add_heading(doc, "4.2 多线程任务", 2)
    p(
        doc,
        "任务按照区县和页码拆分为独立单元，ThreadPoolExecutor 控制并发。每个 future 返回解析条数或错误信息，主线程汇总 total_pages、success_pages、failed_pages、inserted_count、updated_count、unchanged_count 和 snapshot_count。任务最终根据失败情况进入 success、partial_failed 或 failed。",
        "多线程的目的不是压榨目标网站，而是在低并发下减少区县之间的串行等待。实际收口采集使用并发 2、每区 1 页。第一轮 6 个区县全部成功，解析 360 条并新增 360 条；第二轮相同范围有 3 页成功、3 页被网关拦截，仍解析 180 条，其中新增 27 条、未变化 153 条，证明任务能在部分失败时保留有效结果。",
    )
    add_figure(doc, EVIDENCE_DIR / "02-crawl-tasks.png", "图4-1 房天下增量任务与失败日志页面")
    add_heading(doc, "4.3 增量更新实现", 2)
    p(
        doc,
        "入库前先进行字段清洗并计算 fingerprint。若 source + fingerprint 不存在，则插入 listings 并写第一条 snapshot；若存在且价格未变化，则只更新 last_seen_at 和必要的可补全字段；若价格变化，则更新 listings 当前价格并追加 snapshot。该流程由 ListingService.upsert_listing 集中实现，采集器和冷启动导入脚本复用同一规则。",
        f"当前主表 {db['listings_total']:,} 条，快照表 {db['snapshots']:,} 条，恰好多出 1 条，且数据库中有 {db['listings_with_multiple_snapshots']} 套房源存在多次快照。任务 #3 的真实记录为新增 30、价格更新 1、未变化 29、快照新增 1；任务 #6 则证明重复采集 153 条时主表不重复增长。另有 verify_incremental_snapshot.py 可创建临时样本，依次验证插入、重复和价格变化，默认完成后清理。",
    )
    add_code(
        doc,
        """if listing is None:
    insert_listing(cleaned)
    append_snapshot(cleaned)
elif price_changed(listing, cleaned):
    update_current_price(listing, cleaned)
    append_snapshot(cleaned)
else:
    listing.last_seen_at = now""",
        "增量 upsert 核心逻辑（示意）",
    )
    add_heading(doc, "4.4 调度与当前边界", 2)
    p(
        doc,
        "APScheduler 可注册定期质量报告和定期增量采集两个 interval 任务，时区为 Asia/Shanghai，并通过 max_instances=1 防止同一任务并发重入。系统还提供 /api/scheduler/status、/run-quality-report 和 /run-incremental-crawl，用于查看注册状态与手动触发。测试环境和非部署阶段默认关闭常驻调度，属于预期行为。",
        "服务器常驻调度尚未作为已完成上线能力。后续部署时才会设置 SCHEDULER_ENABLED=true、明确区县与页数、通过 systemd 保证进程重启，并观察至少一个完整周期。当前不启用“连续 N 次未出现即 inactive”，因为每次只采有限页面，未出现可能只是本轮未覆盖，贸然标记会引入错误状态。",
    )

    add_heading(doc, "第5章 数据整理（清洗与存储）实现", 1)
    add_heading(doc, "5.1 字段标准化", 2)
    p(
        doc,
        "采集结果先映射到统一 schema，再进入数据库。总价提取数字并统一为万元，单价统一为元/平方米，面积统一为平方米；layout 原文保留并解析 rooms、halls；floor_text 归一为 low、mid、high 或 unknown；build_year 提取年份并派生 house_age；district 使用重庆区县字典映射；link 保留原始链接并参与来源追踪。",
        "异常值不直接删除。总价小于 5 万元或大于 5000 万元、单价小于 1000 或大于 100000 元/平方米、面积小于 10 或大于 500 平方米、建成年份超出合理范围时标记 abnormal。这样既能避免极值污染训练，也能保留原始记录供页面复核和规则调整。",
    )
    add_table(
        doc,
        ["字段", "清洗规则", "异常处理"],
        [
            ["total_price", "提取数字，单位万元", "<5或>5000标记"],
            ["unit_price", "提取数字，单位元/㎡", "<1000或>100000标记"],
            ["area", "提取数字，单位㎡", "<10或>500标记"],
            ["layout", "保留原文并解析室厅", "解析失败保留原文"],
            ["floor", "归一 low/mid/high", "无法判断为 unknown"],
            ["build_year", "提取年份并计算楼龄", "越界标记异常"],
            ["district", "重庆区县标准映射", "失败进入待复核"],
        ],
        [1700, 4100, 3320],
        "表5-1 核心清洗规则",
    )
    add_heading(doc, "5.2 冷启动导入", 2)
    p(
        doc,
        "旧项目中的 SQLite/CSV 只作为一次性冷启动来源。import_legacy_sqlite_to_mysql.py 读取旧表，进行字段映射、清洗和批量 upsert，并把来源标为 anjuke_legacy 或 lianjia_legacy。导入完成后，Web 查询、图表、分析和 Agent 均只访问 MySQL，不再把 SQLite 作为运行数据库。",
        f"当前有效数据中 anjuke_legacy 为 {db['source_counts']['anjuke_legacy']:,} 条，lianjia_legacy 为 {db['source_counts']['lianjia_legacy']:,} 条，fang 为 {db['source_counts']['fang']} 条。该分布说明系统的数据规模主要来自旧库基线，新系统采集用于证明标准链路和增量能力。说明书和答辩必须如实展示这一点，不能把 12 万条都归功于当前爬虫。",
    )
    add_heading(doc, "5.3 数据质量报告", 2)
    p(
        doc,
        f"QualityService 从完整性、极值、低质量样本、可分析记录和快照数量等维度生成报告。最新报告 #{quality['id']} 的 total_count 与业务库一致，为 {quality['total_count']:,} 条；analysis_ready_count 为 {quality['analysis_ready_count']:,} 条；极值记录 {quality['extreme_count']} 条，低质量记录 {quality['low_quality_count']} 条，快照 {quality['snapshot_count']:,} 条。报告已不再停留在旧的 121,837 条。",
        "平均质量分为 100 不代表数据绝对无误，它反映当前规则下必填字段完整度较高。极值检测和模型训练过滤仍然独立执行，前端也展示异常列表。质量分的价值在于形成可复核的运行快照，而不是用单一分数掩盖来源偏差、时间偏差和页面解析误差。",
    )
    add_heading(doc, "5.4 数据规模与分布", 2)
    p(
        doc,
        "区县样本量分布不均衡，沙坪坝、南岸、渝北、江北和九龙坡样本较多，部分远郊区县样本较少。如果直接按更新时间取最近 1000 条，样本会集中在最后导入或最后采集的几个区县，导致模型结论并不代表重庆整体。数据层因此保留区县和来源字段，分析层再按区县进行分层抽样。",
        "趋势图只使用 listings 的 first_seen_at 与 listing_snapshots 的 snapshot_at 等真实时间字段。当前时间跨度主要来自 2025 年旧库和 2026 年系统快照，不能伪造多年趋势。页面中的趋势应理解为现有快照观察，不应外推为长期房地产周期。",
    )

    add_heading(doc, "第6章 数据分析、可视化与智能问数实现", 1)
    add_heading(doc, "6.1 抽样策略修正", 2)
    p(
        doc,
        "早期分析服务按 updated_at 倒序取最近 1000 条，实际只覆盖 4 个区县，虽然运行速度快，却会把局部数据误当作全重庆样本。收口阶段将默认样本扩大到 5000 条，并实现 district_stratified_deterministic：每个有效区县至少保留一个样本，其余名额按区县可用样本量比例分配，区县内部使用固定 CRC32 顺序抽样。",
        f"最新任务 #{job['id']} 的 sample_count 为 {job['sample_count']}，覆盖 {regression['sampling']['district_count']} 个区县和 {regression['sampling']['source_count']} 类来源；其中 anjuke_legacy 4852 条、lianjia_legacy 125 条、fang 23 条。确定性顺序使同一数据集重复训练得到可复核样本，分层覆盖则避免少数区县完全缺失。来源仍以旧库为主，这是模型解释时必须披露的限制。",
    )
    add_heading(doc, "6.2 回归模型", 2)
    p(
        doc,
        "预测目标为挂牌单价，输入特征包含面积、总价辅助字段、房龄、室厅、楼层等级、区县、朝向、装修和来源等。训练前排除强异常和关键字段不足样本，并将类别字段编码。系统比较 RandomForestRegressor、GradientBoostingRegressor、HistGradientBoostingRegressor 等候选模型，以测试集 R² 排序并保存全部候选结果，避免只展示一个模型而没有对照。",
        f"最新最优模型为 {regression['model_name']}。5,000 条源样本中 4,616 条用于回归建模，训练集 {regression['metrics']['train_count']} 条，测试集 {regression['metrics']['test_count']} 条；MAE={regression['metrics']['mae']} 元/平方米，RMSE={regression['metrics']['rmse']} 元/平方米，R²={regression['metrics']['r2']}，MAPE={regression['metrics']['mape']}%。该结果可用于解释特征和提供粗粒度辅助估价，但误差仍然明显，不能称为精准预测。",
        "特征重要性显示房龄贡献最高，区县样本量、楼层等级、朝向、来源和部分区县哑变量也有影响。这里的重要性描述的是模型在当前样本中的分裂贡献，不等于因果关系。房龄重要可能同时反映地段、建筑类型和旧库字段分布，后续需要更均衡的新采集样本和更细的地理特征验证。",
    )
    add_heading(doc, "6.3 KMeans 市场分层", 2)
    p(
        doc,
        f"聚类使用挂牌单价、面积、总价、房龄等连续特征标准化后执行 KMeans，并比较候选 k。最新任务选择 4 类，轮廓系数为 {cluster['metrics']['silhouette_score']}。该值低于早期 4 区县偏置样本的 0.7547，说明覆盖 38 个区县后真实市场结构更连续、类间边界更复杂。指标下降并非系统退化，而是抽样更有代表性后对原先乐观结果的纠正。",
        "聚类结果用于形成样本画像，例如低总价小面积、刚需主流、改善型大面积和高单价核心区等相对分层。每一类都保存样本量、均价、面积和房龄等统计。聚类标签不能直接解释为房屋优劣，也不能替代人工估值；它更适合帮助 Dashboard 筛选和答辩时说明市场内部结构。",
    )
    add_heading(doc, "6.4 异常检测", 2)
    p(
        doc,
        "早期 IsolationForest 在 1000 条样本上识别 369 条异常，36.9%的比例明显过高。修正后污染率按样本规模限制在 2%—8%，当前 5000 条样本取 2%；同时计算区县中位数偏离率与 Z 分数。最终异常由模型孤立样本和强规则共同确认，并为每条记录保存偏离率、Z 分数和命中规则。",
        f"最新结果共识别 {anomaly['metrics']['anomaly_count']} 条异常，占 {anomaly['metrics']['anomaly_rate']*100:.2f}%；其中模型异常 {anomaly['metrics']['model_anomaly_count']} 条，强规则命中 {anomaly['metrics']['strong_rule_count']} 条。这个比例适合人工复核，不会让异常列表淹没正常房源。异常数据默认不物理删除，只在模型训练中过滤。",
    )
    add_heading(doc, "6.5 可视化实现", 2)
    p(
        doc,
        "Dashboard 的 KPI 直接来自 /api/overview，区县排行、价格分布、快照趋势、面积—价格散点和户型分布分别由图表接口返回 JSON。重庆地图加载本地 GeoJSON，支持均价、样本量和质量分切换。房源页支持筛选、分页、详情和 CSV 导出；采集页展示数据源、任务进度与页级日志；质量页展示报告；分析页展示模型证据。",
        "真实页面验证使用 Edge 通道的 CLI Playwright。检查内容包括 URL 与标题、主要 DOM 文本、框架错误覆盖层、console warning/error 和截图。Dashboard 显示 122,316 套、38 个区县和 122,317 条快照；采集页显示 780 条累计解析与部分失败日志；分析页显示 5,000 条样本、38 区县、R² 0.7105 和异常率 2.46%。三页均无相关控制台错误。",
    )
    add_figure(doc, EVIDENCE_DIR / "03-analysis.png", "图6-1 最新分析任务与模型指标页面")
    add_heading(doc, "6.6 DeepSeek Agent", 2)
    p(
        doc,
        "AgentService 先识别问题意图，再由 ToolRegistry 调用 query_market_stats、get_chart_series、get_crawl_status、run_incremental_crawl、run_analysis_job、get_model_result 或 generate_report。只读工具查询服务层聚合结果，写工具只创建受控任务或报告。Agent 不获取数据库连接，不接受任意 SQL，也不能绕过采集规则。",
        f"真实连接测试已改为向 DeepSeek 发起轻量请求，不再仅检查是否配置密钥。首次 max_tokens=8 被推理内容消耗后，将预算调整为 64 并兼容 reasoning_content。数据库现有 {e['agent']['tool_call_count']} 条工具调用和 {e['agent']['generated_report_count']} 份报告，最新报告 #{e['agent']['latest_report']['id']} 同时包含 market 与 model 两类证据；报告 PDF 已导出并通过页数、标题和文本结构检查。",
        "模型问答工具返回前 10 个特征重要性和聚类画像摘要，解决了早期为压缩日志而删掉特征正文的问题。回答中的房源数、均价、MAE、RMSE 和 R² 均来自工具 JSON。若工具没有数据，系统提示先执行采集或分析任务，而不是编造具体数值。",
    )
    add_figure(doc, EVIDENCE_DIR / "06-agent-evidence.png", "图6-2 Agent 工具调用与报告验收脚本输出快照")

    add_heading(doc, "第7章 测试、问题解决、总结与展望", 1)
    add_heading(doc, "7.1 测试与性能", 2)
    p(
        doc,
        "后端测试覆盖健康检查、鉴权、房源筛选、Dashboard、采集、质量、调度、分析、Agent 和系统设置。收口后执行 python -m compileall Backend 成功，python -m pytest Backend/tests -q 返回 34 passed；前端 npm run build 成功，Vite 转换 2944 个模块。构建产物主 JavaScript 约 2.1 MB，存在分包建议，但不影响本地课程演示，页面懒加载列为后续优化。",
        "常用图表接口在真实 12 万条业务库上进行一次本机冷调用测试，overview 为 1019.58 ms，district-price 为 253.42 ms，price-distribution 为 306.67 ms，price-trend 为 98.10 ms，500 点散点为 25.92 ms，户型分布为 162.93 ms，均低于 2 秒目标。该结果只代表当前电脑与数据库状态，不宣称等同于公网部署性能。",
    )
    add_table(
        doc,
        ["验收项", "命令/路径", "结果"],
        [
            ["语法检查", "python -m compileall Backend", "通过"],
            ["后端测试", "python -m pytest Backend/tests -q", "34 passed"],
            ["前端构建", "npm run build", "通过，2.1 MB提示"],
            ["增量验证", "python scripts/verify_incremental_snapshot.py", "脚本可复用"],
            ["接口性能", "/api/overview 与 /api/charts/*", "全部 < 2 秒"],
            ["浏览器验证", "CLI Playwright + Edge", "三核心页通过"],
        ],
        [2000, 4300, 2820],
        "表7-1 非部署阶段验收结果",
    )
    add_heading(doc, "7.2 关键问题与解决过程", 2)
    p(
        doc,
        "问题一是模型样本偏差。表面上 1000 条足够训练，实际按更新时间取样只覆盖 4 个区县，导致较高的聚类轮廓系数和不可靠的全市结论。解决方法不是简单扩大 limit，而是按区县分层分配名额，并记录抽样策略、区县覆盖和来源分布。修正后回归 R² 提升到 0.7105，但聚类轮廓系数降到 0.2848，更符合全市数据的复杂性。",
        "问题二是异常比例过高。单独使用 IsolationForest 且污染率设置不合理，会把大量正常波动列为异常。解决方法是收紧污染率、引入区县中位数偏离和 Z 分数，并输出判定理由。异常率降到 2.46%，更适合复核。问题三是外部页面波动。系统不重复强试，而是把验证码或网关拦截记录为 WARN，使任务部分失败但不丢失成功页。",
        "问题四是 Agent 连接测试过于形式化。只判断密钥存在不能证明接口可用，因此改为真实轻量请求；又针对推理模型可能把少量 token 用于 reasoning_content 的情况提高预算并兼容诊断预览。问题五是文档与实现脱节，旧稿仍写原生 JS、未来功能和部署完成态。本次以业务库、测试和页面证据为准统一重写，部署只作为后续方案。",
    )
    add_heading(doc, "7.3 专业能力与项目管理总结", 2)
    p(
        doc,
        "本项目训练了从需求拆解到验证交付的完整能力。开发过程中需要同时理解网页解析、数据建模、SQL 聚合、后端分层、前端状态、机器学习指标和大模型工具调用。相比只完成单个算法实验，系统化工作要求每个数字都能追溯到数据库或模型结果，每个“已完成”都要有命令、接口或页面证据。",
        "在项目管理上，采用 GOAL 任务标准把目标、产物、验收和边界写清楚，避免一次性重构整个系统。遇到浏览器、网页和模型接口问题时遵循失败次数与退出规则，不在同一路径反复消耗。对旧数据、新采集、模型边界和部署状态保持如实表述，是课程设计中比堆砌技术名词更重要的工程诚信。",
        "团队协作方面，本文只记录本人张浩博实际承担并可核验的需求梳理、数据链路、分析建模、前后端联调、测试和材料收口工作，不虚构未确认成员的分工。若最终以小组提交，可由组长依据真实提交记录补充其他成员的任务和贡献。",
    )
    add_heading(doc, "7.4 当前不足", 2)
    p(
        doc,
        "第一，新系统标准采集数据只有 539 条，模型样本仍以旧库为主；需要长期低频采集提高新数据占比。第二，时间快照跨度有限，趋势分析只能描述现有采集窗口。第三，质量分对完整性较敏感，但对语义错误、重复小区名称和平台偏差的度量仍较粗。第四，KMeans 轮廓系数仅 0.2848，说明市场分层边界不强，不能把聚类包装成确定分类。",
        "第五，前端主包约 2.1 MB，后续可按路由懒加载图表。第六，Agent 初始页面不自动恢复数据库中的历史会话，真实报告需要在当前会话生成后展示；数据库证据和 PDF 导出已完成，但会话持久化 UI 可继续完善。第七，调度器尚未在服务器连续运行，部署、HTTPS、日志轮转、备份恢复和公网性能均未验收。",
    )
    add_heading(doc, "7.5 可复现性与工程细节复盘", 2)
    add_heading(doc, "7.5.1 查询设计与索引", 3)
    p(
        doc,
        "房源主表的数据量达到十万级后，页面性能不再只取决于前端渲染，数据库筛选和聚合方式更加关键。listings 对 district、total_price 和 updated_at 建立索引，source 与 fingerprint 使用联合唯一约束。房源分页查询先组合区县、价格、面积和关键词条件，再执行 count 与分页结果查询；图表接口只返回绘图需要的聚合字段或有限散点，不把整表序列化给浏览器。",
        "overview 需要同时统计有效房源、均价、总价、完整率、区县数和快照数，冷调用约 1 秒，是常用接口中耗时最高的一项。区县排行、价格分布和户型分布使用 GROUP BY，散点图通过 limit 控制返回量，趋势图只聚合快照月份。当前指标已经满足课程演示，但如果未来数据持续增长，可以增加按日预聚合表、短时缓存和更细的联合索引，并用 EXPLAIN 验证执行计划，而不是过早引入复杂中间件。",
    )
    add_heading(doc, "7.5.2 配置、鉴权与秘密信息", 3)
    p(
        doc,
        "配置分为代码默认值、.env 环境变量和 system_settings 数据库设置三层。开发环境可以使用默认端口，但数据库密码、SECRET_KEY 和 DeepSeek API Key 不应写进源码。设置接口返回时只显示掩码，更新密钥时识别占位掩码并避免覆盖真实值。测试环境强制关闭 DeepSeek 与调度器，防止自动化测试产生外部费用或修改业务任务。",
        "登录模块定位为课程演示的本地管理员鉴权，Token 有过期时间，前端 API 客户端统一附加 Authorization 头。系统没有实现学校 SSO、找回密码、用户注册和多角色权限，因此登录页已经移除相关入口，文档也不把它描述为企业级权限系统。若后续公网部署，应更换默认账号密码、启用 HTTPS、限制跨域来源，并将鉴权状态写入安全日志。",
    )
    add_heading(doc, "7.5.3 测试隔离与故障注入", 3)
    p(
        doc,
        "后端测试使用 real_estate_test，conftest 在测试应用中关闭鉴权、调度器和 DeepSeek，保证测试不会污染 122,316 条业务数据。采集器测试通过本地 HTML 或模拟响应验证解析，任务测试检查成功、失败和部分失败状态；分析测试使用小样本验证分层覆盖、指标字段与异常率上限；Agent 测试验证工具选择、证据裁剪和特征重要性保留。",
        "真实故障并不都适合在单元测试中模拟，因此收口阶段还保留运行证据。房天下第二轮出现三页网关拦截，任务仍保存三页成功结果并以 partial_failed 结束；Playwright 默认浏览器缺失时切换到已安装 Edge，页面引用失效或交互没有推进时按退出规则停止，而不是在多个浏览器工具之间循环。这些记录体现了系统对外部不确定性的处理方式。",
    )
    add_heading(doc, "7.5.4 模型复现与证据字段", 3)
    p(
        doc,
        "模型结果不仅保存 MAE、RMSE、R² 等指标，还保存抽样策略、请求样本上限、实际样本数、区县分布、来源分布和固定种子。回归候选模型保存排名与测试指标，聚类保存候选 k 和画像，异常检测保存污染率、阈值、模型异常数与强规则数。这样即使页面只展示摘要，也可以从 model_results.evidence_json 还原结论条件。",
        "固定 CRC32 顺序抽样的价值是让同一数据版本能够重复得到相同样本，便于比较代码改动前后的结果；它不意味着样本天然随机，也不能消除旧库来源偏差。为了进一步提高科研严谨性，后续可以保存数据版本哈希、训练参数、scikit-learn 版本和模型文件，并采用按区县分组的交叉验证，检查模型是否只记住样本量大的区域。",
    )
    add_heading(doc, "7.5.5 文档证据自动化与数据伦理", 3)
    p(
        doc,
        "正式材料中的数字来自 export_acceptance_evidence.py 导出的 JSON，而不是手工复制旧结论。脚本读取业务库的房源、来源、快照、质量、模型、采集和 Agent 表，并执行常用接口耗时测试；build_submission_documents.py 再使用学院模板生成说明书、安装说明、答辩清单和日志。数据变化后可重新运行两条命令刷新材料，减少文档与系统不一致。",
        "自动生成不等于可以省略人工审查。文档明确区分挂牌价与成交价、旧库与新采集、代码支持与服务器常驻、模型相关性与因果解释。网页采集遵守低并发和停止规则，验证码或网关出现时记录失败，不尝试绕过；房源数据用于课程研究，不在文档中展示个人联系方式。上述边界既是技术限制，也是数据工程必须承担的伦理责任。",
    )
    add_heading(doc, "7.5.6 前端状态与可用性设计", 3)
    p(
        doc,
        "前端页面不是把接口 JSON 直接堆在屏幕上，而是针对不同任务设计状态。Dashboard 区分加载、成功和刷新；房源表在筛选变化时回到第一页，避免页码超出结果范围；采集任务用进度条、状态标签和日志等级区分正常完成、部分失败与失败；分析页只有在最新任务成功且结果齐全时显示指标，缺少结果时给出重新训练入口。",
        "页面视觉采用深蓝导航、白色卡片和浅灰背景，数值单位与标签分离，避免把挂牌单价和总价混在同一尺度。图表标题说明数据口径，地图支持均价、样本量和质量切换。浏览器验证重点检查 1440×1000 桌面答辩视口中的对齐、溢出、图表渲染、favicon 和控制台错误；移动端不是本轮主要演示场景，因此只保留响应式基础能力，不宣称完成全面移动端适配。",
    )
    add_heading(doc, "7.5.7 数据血缘与口径核验", 3)
    p(
        doc,
        "每条房源通过 source、source_listing_id、link、first_seen_at 和 last_seen_at 保留基本血缘。冷启动导入脚本不会把 legacy 来源改写成 fang；采集任务的 task_id 可以关联快照；质量报告保存生成时间；模型结果关联 analysis_job；Agent 工具调用保存 session_id、question、tool_name、args、result 和 duration。由此可以从页面结论逐层追溯到服务结果和数据库记录。",
        "口径核验尤其关注区县名称和时间字段。旧库中可能出现带“区”与不带“区”的名称，系统在分析证据中保留真实分布，前端展示时做可读映射但不伪造合并结果。快照趋势按真实月份聚合，若某个月没有采集就不补零生成虚假曲线。模型目标始终是 unit_price，报告中的平均总价与平均单价分别注明万元和元/平方米。",
    )
    add_heading(doc, "7.5.8 答辩可解释性设计", 3)
    p(
        doc,
        "答辩演示按照“问题—系统—证据—边界”组织，而不是逐页介绍技术名词。先说明公开挂牌数据存在来源分散、重复和变化问题，再展示 MySQL 规模与来源结构；随后用采集任务证明失败可追踪、重复不重复入库和价格变化进快照；再用 Dashboard 与模型页说明分析结论；最后用 Agent 工具记录说明智能问数没有绕开数据治理。",
        "对指标的解释坚持保守口径。R² 0.7105 表明模型解释了测试集约 71%的波动，但 MAE 仍有 1166.21 元/平方米；轮廓系数 0.2848 表明市场分层只是辅助画像；异常率 2.46% 表示待复核样本，不表示错误房源；12万条规模主要来自旧库冷启动，新系统采集为539条。把这些限制主动讲清楚，比只展示最好看的数字更能体现数据分析的可信度。",
    )
    add_heading(doc, "7.5.9 维护、备份与版本管理", 3)
    p(
        doc,
        "运行维护需要同时保护代码、配置和数据。代码通过 Git 记录版本，.env、数据库密码、API Key、运行日志和临时截图不应提交；数据库应定期使用 mysqldump 生成带日期的备份，并至少进行一次恢复演练。采集、质量和分析任务都保留 created_at、started_at 与 finished_at，便于发现卡死任务和比较不同版本的运行结果。",
        "本轮只在本机完成业务库验证，没有执行生产备份和恢复，因此文档不把备份脚本写成已上线能力。后续部署时应增加日志轮转、磁盘空间告警、数据库最小权限账号和备份保留周期；模型与报告需要关联代码版本和数据版本。只有备份文件能够实际恢复、服务重启能够自动拉起，维护方案才算完成，而不是仅存在一段部署命令。",
    )
    add_heading(doc, "7.6 后续部署方案", 2)
    p(
        doc,
        "部署阶段计划使用阿里云轻量服务器或 ECS，MySQL 8.x 作为数据库，Gunicorn 承载 Flask，Nginx 提供静态文件和反向代理，systemd 管理进程，环境变量保存数据库连接与 DeepSeek 密钥。上线后先导入冷启动数据，再执行小规模健康检查与增量任务，最后开启质量报告和增量采集调度。",
        "部署验收必须包含公网首页和 API 可访问、服务器重启后服务自动恢复、Nginx 与 Gunicorn 日志正常、数据库备份可恢复、敏感信息未提交。由于本轮明确暂缓部署，以上内容仅为方案，不作为已经上线的事实。",
    )
    add_heading(doc, "7.7 结论", 2)
    p(
        doc,
        "项目已经完成学院要求的核心非部署闭环：超过 5 万条的 MySQL 数据底座、字段清洗、增量 upsert 与快照、多线程任务和失败日志、Web 多维可视化、回归聚类异常检测、DeepSeek 工具调用与报告、自动化测试和正式材料。最重要的收口改进是把模型样本扩展为覆盖 38 个区县的分层样本，并把数据来源、模型误差、异常比例和未完成部署如实写入证据。",
        "系统当前适合作为课程设计原型和后续科研工程基础，而不是商业房价估值产品。只要继续积累新系统快照、完善会话持久化、优化前端分包并完成服务器部署，就可以进一步形成可长期运行、可公开演示的重庆二手房挂牌价智能分析平台。",
    )

    body_text = "".join(par.text for par in doc.paragraphs if par.text and "目录" not in par.text)
    body_text += "".join(cell.text for table in doc.tables[1:] for row in table.rows for cell in row.cells)
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", body_text))
    if chinese_chars < 10000:
        raise RuntimeError(f"正文中文字符不足：{chinese_chars}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out = OUTPUT_DIR / f"{TITLE}_学年设计说明书.docx"
    doc.core_properties.title = TITLE + " 学年设计说明书"
    doc.core_properties.author = STUDENT
    doc.core_properties.subject = "学年设计Ⅱ正式提交文档"
    doc.core_properties.comments = f"正文中文字符约 {chinese_chars}；部署状态：暂缓。"
    doc.save(out)
    print(f"[ok] main document: {out} chinese_chars={chinese_chars}")
    return out


def new_aux(title: str, subtitle: str) -> Document:
    doc = Document()
    configure_styles(doc, compact=True)
    for section in doc.sections:
        set_page(section)
    configure_header_footer(doc, TITLE + "｜" + title)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=22, bold=True, color=ACCENT)
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(18)
    r = sp.add_run(subtitle)
    set_run_font(r, size=10, color="666666")
    return doc


def build_install(e: dict) -> Path:
    doc = new_aux("系统安装与使用说明", "非部署阶段本地运行版｜2026年6月")
    add_heading(doc, "1. 适用范围", 1)
    p(doc, "本说明用于在 Windows 本地环境启动重庆二手房挂牌价分析系统。运行期数据库必须为 MySQL 8.x；旧 SQLite/CSV 只用于一次性冷启动导入。部署到阿里云不在本轮范围内。")
    add_heading(doc, "2. 环境要求", 1)
    add_table(doc, ["组件", "建议版本", "用途"], [["Python", "3.10+", "Flask后端、采集与分析"], ["MySQL", "8.x", "唯一运行数据库"], ["Node.js/npm", "20+", "Vite前端"], ["浏览器", "Edge/Chrome", "页面演示"], ["DeepSeek Key", "可选", "真实Agent问答"]], [1800, 1800, 5520])
    add_heading(doc, "3. 初始化步骤", 1)
    for item in ["在 MySQL 中创建 real_estate 与 real_estate_test 数据库。", "复制 .env.example 为 .env，填写 DATABASE_URL、TEST_DATABASE_URL 和管理员配置。", "在项目根目录执行 pip install -r requirements.txt。", "在 Frontend 目录执行 npm install。", "如需导入旧库，运行 python scripts/import_legacy_sqlite_to_mysql.py 并核对来源统计。"]:
        add_bullet(doc, item)
    add_code(doc, """DATABASE_URL=mysql+pymysql://user:password@127.0.0.1:3306/real_estate?charset=utf8mb4
TEST_DATABASE_URL=mysql+pymysql://user:password@127.0.0.1:3306/real_estate_test?charset=utf8mb4
AUTH_REQUIRED=true
DEEPSEEK_ENABLED=true
DEEPSEEK_API_KEY=请在本地填写，禁止提交仓库
SCHEDULER_ENABLED=false""", ".env 关键配置示例")
    add_heading(doc, "4. 启动方式", 1)
    p(doc, "推荐在项目根目录运行 PowerShell 脚本，它会以隐藏窗口启动 Flask 与 Vite，并把日志写入 .codex-run。也可以分别手动启动后端和前端。")
    add_code(doc, """powershell -ExecutionPolicy Bypass -File scripts/start_local_dev.ps1

# 手动后端
python -m flask --app Backend.app run --debug --port 5000

# 手动前端（另开终端）
cd Frontend
npm run dev""")
    add_heading(doc, "5. 使用流程", 1)
    for item in ["访问 http://127.0.0.1:5173，使用本地管理员账号登录。", "首页总览查看房源数、区县覆盖、快照和图表。", "房源数据管理进行区县、价格、面积和关键词筛选。", "采集任务管理只创建小规模低并发任务，并查看失败日志。", "分析建模可查看最新任务或重新运行 all 分析。", "智能问答只在配置 DeepSeek 后使用，所有数值应有工具调用证据。"]:
        add_bullet(doc, item)
    add_heading(doc, "6. 验收命令", 1)
    add_code(doc, """python -m compileall Backend
python -m pytest Backend/tests -q
cd Frontend
npm run build

# 根目录执行
python scripts/verify_incremental_snapshot.py
python scripts/export_acceptance_evidence.py
python scripts/local_demo_smoke.py --base-url http://127.0.0.1:5000""")
    add_heading(doc, "7. 常见问题", 1)
    add_table(doc, ["现象", "处理"], [["数据库连接失败", "核对MySQL进程、端口、库名和DATABASE_URL"], ["前端接口失败", "确认5000端口后端运行，检查VITE_API_BASE_URL"], ["采集部分失败", "查看crawl_logs；遇验证码停止重试，不做绕过"], ["Agent不可用", "在设置页执行真实连接测试，检查模型与密钥"], ["调度器未运行", "本地默认关闭；部署前再开启SCHEDULER_ENABLED"]], [3000, 6120])
    add_heading(doc, "8. 数据口径", 1)
    p(doc, f"当前业务库有效房源 {e['database']['valid_listings']:,} 条，其中新系统 fang 标准数据 {e['database']['source_counts']['fang']} 条，其余主要为旧库冷启动基线。所有价格均为挂牌价/报价。")
    out = OUTPUT_DIR / f"{TITLE}_安装与使用说明.docx"
    doc.save(out)
    print(f"[ok] install guide: {out}")
    return out


def build_evidence(e: dict) -> Path:
    doc = new_aux("答辩证据清单", "证据日期：2026年6月18日｜部署项明确暂缓")
    add_heading(doc, "1. 核心结论", 1)
    p(doc, f"业务库有 {e['database']['valid_listings']:,} 条有效房源、38个区县、{e['database']['snapshots']:,} 条快照；最新模型样本5000条，R²={e['model_results']['regression']['metrics']['r2']}；Agent保存{e['agent']['tool_call_count']}条工具调用和{e['agent']['generated_report_count']}份报告。")
    add_heading(doc, "2. 逐项证据", 1)
    rows = [
        ["1", "5万条与区县覆盖", "05-database-evidence.png；Dashboard", "已具备"],
        ["2", "来源与冷启动说明", "acceptance_evidence.json/source_counts", "已具备"],
        ["3", "采集成功与失败日志", "02-crawl-tasks.png；crawl_tasks #5/#6", "已具备"],
        ["4", "重复采集不重复入库", "任务#6：未变153、新增27", "已具备"],
        ["5", "价格变化进入快照", "任务#3：updated=1、snapshot=1", "已具备"],
        ["6", "质量报告最新", f"data_quality_reports #{e['quality']['id']}", "已具备"],
        ["7", "Dashboard可视化", "01-dashboard.png", "已具备"],
        ["8", "模型指标与覆盖", "03-analysis.png；任务#9", "已具备"],
        ["9", "Agent工具与PDF报告", "06-agent-evidence.png；报告#5 PDF", "已具备"],
        ["10", "阿里云公网部署", "Nginx/Gunicorn/systemd/MySQL", "本轮暂缓"],
    ]
    add_table(doc, ["序号", "验收点", "证据位置", "状态"], rows, [800, 2500, 4000, 1820])
    add_heading(doc, "3. 答辩演示顺序", 1)
    for item in ["先说明12万数据主要是旧库冷启动，不回避来源结构。", "打开Dashboard展示房源数、38区县、快照和区县图表。", "进入采集任务展示360条新增、153条未变和网关拦截日志。", "说明fingerprint不含价格，展示一次真实价格快照。", "进入分析页展示5000条分层样本、R²、轮廓系数和异常率。", "展示Agent工具记录与报告#5 PDF，强调数值来自工具JSON。", "最后说明部署暂缓和后续方案。"]:
        add_bullet(doc, item)
    add_heading(doc, "4. 推荐回答口径", 1)
    add_table(doc, ["问题", "回答要点"], [["12万条都是新爬的吗？", "不是，121,777条为旧库冷启动，539条为新系统fang标准数据。"], ["为什么聚类指标下降？", "分层抽样覆盖38区县后更真实，早期4区县结果偏乐观。"], ["能预测成交价吗？", "不能，目标是挂牌单价辅助估计，MAE约1166元/㎡。"], ["为什么调度器没常驻？", "非部署阶段默认关闭，注册、接口和手动增量已验证。"], ["Agent会编数据吗？", "具体数值必须来自白名单工具JSON，调用记录持久化。"]], [3000, 6120])
    out = OUTPUT_DIR / f"{TITLE}_答辩证据清单.docx"
    doc.save(out)
    print(f"[ok] evidence checklist: {out}")
    return out


def build_log(e: dict) -> Path:
    doc = new_aux("学年设计日志", f"学生：{STUDENT}｜项目：{TITLE}")
    add_heading(doc, "1. 日志说明", 1)
    p(doc, "日志根据仓库文件、数据库任务时间和本轮实际验收记录整理，未虚构部署、长期调度或未确认组员工作。")
    rows = [
        ["2026-06-08", "读取指导手册与模板，整理5万条、增量、可视化、挖掘和8000字文档要求。", "形成需求边界与项目资料目录。"],
        ["2026-06-10", "完成Flask/MySQL/React核心模块核对，检查房源、任务、Dashboard、模型和Agent结构。", "明确旧库冷启动与新系统采集口径。"],
        ["2026-06-11", "执行真实页面与接口检查，补充采集、快照、系统设置和PDF导出证据。", "确认主链路可运行。"],
        ["2026-06-12", "补充本地启动、smoke check与增量快照脚本，完善README。", "形成可复用验收命令。"],
        ["2026-06-18 上午", "将分析改为5000条区县分层样本，收紧异常检测，重跑真实模型。", "38区县；R²=0.7105；异常率2.46%。"],
        ["2026-06-18 上午", "房天下6区低并发补采与重复采集，刷新质量报告。", "fang=539；重复153条未变；失败页可追踪。"],
        ["2026-06-18 上午", "执行DeepSeek真实连接、问数、模型解释和报告生成。", f"工具调用{e['agent']['tool_call_count']}条；报告{e['agent']['generated_report_count']}份。"],
        ["2026-06-18 中午", "修正前端技术文案、favicon和分析证据字段，执行构建与Playwright截图。", "核心页面无相关控制台错误。"],
        ["2026-06-18 下午", "重做说明书、安装说明、证据清单和设计日志，统一非部署口径。", "形成正式提交材料。"],
    ]
    add_table(doc, ["日期", "工作内容", "结果/问题解决"], rows, [1600, 4900, 2620])
    add_heading(doc, "2. 个人总结", 1)
    p(doc, "本次学年设计的主要收获是建立了证据驱动的开发习惯：数据规模要能用SQL核验，模型结论要有样本覆盖和误差指标，页面完成要有真实后端数据和控制台检查，Agent回答要有工具调用记录。遇到网页和浏览器自动化问题时遵守失败退出规则，比反复尝试更能保证进度。")
    out = OUTPUT_DIR / f"{TITLE}_学年设计日志.docx"
    doc.save(out)
    print(f"[ok] design log: {out}")
    return out


def main() -> int:
    parser = argparse.ArgumentParser(description="基于学院模板和验收JSON生成正式提交材料。")
    parser.add_argument("--evidence", type=Path, default=EVIDENCE)
    args = parser.parse_args()
    if not TEMPLATE.exists():
        raise FileNotFoundError(TEMPLATE)
    if not args.evidence.exists():
        raise FileNotFoundError(args.evidence)
    evidence = json.loads(args.evidence.read_text(encoding="utf-8"))
    create_evidence_cards(evidence)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    build_main(evidence)
    build_install(evidence)
    build_evidence(evidence)
    build_log(evidence)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
