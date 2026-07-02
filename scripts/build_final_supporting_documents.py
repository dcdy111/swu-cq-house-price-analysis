from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "deliverables"
TITLE = "重庆市二手房源价格数据分析与智能可视化系统"
STUDENT = "张浩博"
DOC_DATE = "2026年7月1日"

NAVY = "173F73"
BLUE = "315F91"
PALE_BLUE = "EAF0F8"
PALE_GREEN = "E8F4E8"
PALE_YELLOW = "FFF7DD"
LIGHT_GRAY = "F3F4F6"
MID_GRAY = "D8DEE8"
TEXT = "1F2937"
MUTED = "64748B"
WHITE = "FFFFFF"

# A4 institutional override for the compact-reference visual system.
PAGE_WIDTH_CM = 21.0
PAGE_HEIGHT_CM = 29.7
MARGIN_CM = 2.4
CONTENT_WIDTH_DXA = 9184
TABLE_INDENT_DXA = 120


def set_run_font(run, east_asia="宋体", latin="Times New Roman", size=10.5, bold=None, color=TEXT):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), east_asia)
    rpr.rFonts.set(qn("w:ascii"), latin)
    rpr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent_dxa=TABLE_INDENT_DXA):
    if sum(widths) != CONTENT_WIDTH_DXA:
        raise ValueError(f"table widths must sum to {CONTENT_WIDTH_DXA}: {widths}")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_borders(table, color=MID_GRAY, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document, running_label: str):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.3
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)

    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 9, NAVY),
        ("Heading 2", 13, 13, 6, BLUE),
        ("Heading 3", 11.5, 9, 4, TEXT),
    ):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(10.5)
        style.paragraph_format.left_indent = Cm(0.9)
        style.paragraph_format.first_line_indent = Cm(-0.45)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.space_after = Pt(4)

    for section in doc.sections:
        section.page_width = Cm(PAGE_WIDTH_CM)
        section.page_height = Cm(PAGE_HEIGHT_CM)
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)
        section.header_distance = Cm(1.1)
        section.footer_distance = Cm(1.1)

        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hp.paragraph_format.space_after = Pt(0)
        hr = hp.add_run(running_label)
        set_run_font(hr, east_asia="微软雅黑", size=8.5, color=MUTED)

        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(0)
        r = fp.add_run("第 ")
        set_run_font(r, size=9, color=MUTED)
        add_page_number(fp)
        r = fp.add_run(" 页")
        set_run_font(r, size=9, color=MUTED)


def add_cover(doc: Document, title: str, subtitle: str, metadata: list[tuple[str, str]]):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(44)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体", size=24, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run(subtitle)
    set_run_font(r, east_asia="微软雅黑", size=12, color=MUTED)

    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_geometry(table, [2050, 7134])
    set_table_borders(table, color=WHITE, size=0)
    for index, (label, value) in enumerate(metadata):
        left, right = table.rows[index].cells
        shade_cell(left, PALE_BLUE)
        shade_cell(right, "F8FAFC")
        lp = left.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        lr = lp.add_run(label)
        set_run_font(lr, east_asia="微软雅黑", size=10, bold=True, color=NAVY)
        rp = right.paragraphs[0]
        rr = rp.add_run(value)
        set_run_font(rr, size=10, color=TEXT)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(28)
    r = p.add_run("西南大学商贸学院")
    set_run_font(r, east_asia="微软雅黑", size=11, bold=True, color=NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DOC_DATE)
    set_run_font(r, size=10, color=MUTED)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level=1):
    return doc.add_paragraph(text, style=f"Heading {level}")


def add_body(doc: Document, text: str, first_indent=True):
    p = doc.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Cm(0.74) if first_indent else Cm(0)
    r = p.add_run(text)
    set_run_font(r)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        set_run_font(r)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade_cell(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, east_asia="微软雅黑", size=9.2, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        for index, value in enumerate(values):
            cell = row.cells[index]
            if row_index % 2 == 1:
                shade_cell(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index == 0 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=TEXT)
    # Newly appended rows need the same fixed geometry as the header row.
    set_table_geometry(table, widths)
    tail = doc.add_paragraph()
    tail.paragraph_format.space_after = Pt(0)
    return table


def add_code(doc: Document, code: str, caption: str | None = None):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(caption)
        set_run_font(r, east_asia="微软雅黑", size=9, bold=True, color=BLUE)
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=MID_GRAY)
    cell = table.cell(0, 0)
    shade_cell(cell, "F7F9FC")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(code.strip().splitlines()):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, east_asia="等线", latin="Consolas", size=8.5, color="243247")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_note(doc: Document, label: str, text: str, fill=PALE_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color="E7D79A")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + "：")
    set_run_font(r, east_asia="微软雅黑", size=9.5, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r, size=9.5, color=TEXT)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build_log() -> Path:
    doc = Document()
    configure_document(doc, TITLE + "｜学年设计日志")
    add_cover(
        doc,
        "学年设计日志",
        "基于 Git、数据库任务记录和实际验收证据整理",
        [
            ("项目名称", TITLE),
            ("学生", STUDENT),
            ("课程", "学年设计Ⅱ"),
            ("记录区间", "2026年6月10日—2026年7月1日"),
            ("事实口径", "只记录已实现、已运行或有提交记录的工作"),
        ],
    )

    add_heading(doc, "1. 日志编制说明", 1)
    add_body(
        doc,
        "本日志依据仓库 Git 提交历史、MySQL 中的采集任务与分析任务日期、自动化测试结果和真实页面验收记录整理。日志不把规划写成已完成，不虚构服务器部署、团队成员姓名或未发生的后续日期。课程正式动员前的记录统一标注为前期准备。",
    )
    add_note(
        doc,
        "当前验收基线",
        "MySQL 有效房源 100,897 条，Dashboard 标准化展示 38 个区县，来源包括 fang 98,134 条、anjuke_mobile 2,753 条、lianjia 10 条；快照 100,914 条。",
        fill=PALE_BLUE,
    )

    add_heading(doc, "2. 过程记录", 1)
    rows = [
        ["2026-06-10", "前期准备", "建立项目初始快照，梳理 Flask、MySQL、前端和数据分析的总体目录。", "Git：f0977c8；形成可追踪工程基线。"],
        ["2026-06-11", "核心闭环", "完成采集、存储、Dashboard、分析和 Agent 的基础交付链，并补充鉴权、真实数据与 PDF 导出。", "Git：d3f871b、2525bf2；核心模块可联调。"],
        ["2026-06-11", "异常处理", "针对房天下验证页面增加识别和失败记录，避免把验证页误解析为房源。", "Git：9789ba7；采集失败可记录而不导致服务崩溃。"],
        ["2026-06-20", "版本整理", "整理重庆二手房分析系统初始版本，统一项目入口和主要功能页面。", "Git：0871569；形成可持续迭代版本。"],
        ["2026-06-26", "采集稳定性", "完善实时采集任务、低并发控制、页级日志、增量 upsert 和快照证据。", "Git：23505e3；数据库记录当日采集任务 20 个。"],
        ["2026-06-26", "前端验收", "修复页面联调流程，核对登录、Dashboard、房源筛选、任务和分析页面。", "Git：cf2b6a0；核心页面可使用真实接口数据。"],
        ["2026-06-30", "需求复核", "阅读指导手册和课程说明会记录，确认 5 万条、增量维护、Web 可视化、挖掘结论、安装说明和日志要求。", "形成老师要求与系统证据对应表。"],
        ["2026-06-30", "地图与透明性", "接入高德 JS API 地图并保留本地重庆 GeoJSON 数据层，补充数据来源透明说明。", "Git：e91e2d3；地图支持均价、样本量和质量切换。"],
        ["2026-06-30", "界面优化", "将系统调整为浅色科研后台风格，优化 Dashboard、分析页和数据表的 PC 端可读性。", "Git：1b97831、050427c；答辩视口布局更稳定。"],
        ["2026-06-30", "Agent 完善", "完善 DeepSeek 会话、流式回答、工具调用记录、报告与前端执行轨迹展示。", "Git：7599a4e；数据库保留工具调用和报告证据。"],
        ["2026-07-01", "数据收口", "完成最终 MySQL 数据核验与 SQL 导出，检查来源、区县、快照、任务和质量报告。", "100,897 条有效房源；100,914 条快照；3 个真实来源。"],
        ["2026-07-01", "分析建模", "运行 EDA、回归、KMeans 和 IsolationForest 分析任务，保存候选模型和指标。", "任务 #20 成功；GBDT：MAE 1991.61、RMSE 2673.64、R² 0.3847。"],
        ["2026-07-01", "自然语言查数", "增加 DeepSeek 生成只读 MySQL SELECT 的白名单工具，加入 AST 校验、超时和行数限制。", "Git：f0fb950；真实问数调用 query_readonly_sql 成功。"],
        ["2026-07-01", "图表与布局", "修复总价分布、户型图、空图表和分析页左下大块空白，重点验证 PC 端。", "1920×900 页面复验通过，控制台无相关错误。"],
        ["2026-07-01", "运行环境恢复", "恢复本地 .env、前端 .env.local、依赖与 Cookie 配置，并用 .gitignore 隔离敏感信息。", "真实 DeepSeek 连接成功；本地配置保留但不进入 Git。"],
        ["2026-07-01", "交付脚本", "新增不覆盖现有配置的一键安装脚本，修复启动脚本端口状态误判和 Vite 启动等待。", "Git：c3d9952；5000/5173 停启复验通过。"],
        ["2026-07-01", "最终测试", "执行后端全量测试、前端生产构建、接口 smoke 和真实浏览器检查。", "82 passed；npm run build 通过；核心接口 smoke 全部通过。"],
    ]
    add_table(doc, ["日期", "阶段", "完成工作", "结果与证据"], rows, [1260, 1200, 3920, 2804], font_size=8.5)

    add_heading(doc, "3. 关键问题与解决记录", 1)
    issue_rows = [
        ["采集页面出现验证码或验证网关", "把验证页面识别为失败证据，限制重试次数，不绕过验证码。", "任务可部分失败，成功页和失败页均保留日志。"],
        ["重复采集可能造成主表膨胀", "使用 source + fingerprint 定位房源；fingerprint 不包含价格。", "重复记录只更新 last_seen；价格变化才追加快照。"],
        ["图表标签重叠和分析页留白", "调整图表方向、图例布局、空状态和历史任务网格。", "PC 页面无明显遮挡，左下空白问题消除。"],
        ["自然语言 SQL 存在安全风险", "只接受模型基于自然语言生成的 SELECT，执行 AST、表白名单、LIMIT、只读事务和超时校验。", "用户原始 SQL 不直接执行，写操作和系统表被拒绝。"],
        ["本地配置被误清理", "从回收站恢复 .env 和 .env.local，保留依赖与运行产物，仅用 .gitignore 隔离。", "本地全功能恢复，真实 DeepSeek 和调度器均可用。"],
        ["启动脚本误判端口", "端口检测限定 Listen 状态，并把 Vite 固定等待改为最多 20 秒轮询。", "连续停启后健康检查和前端 200 均通过。"],
    ]
    add_table(doc, ["问题", "处理方法", "验证结果"], issue_rows, [2500, 3880, 2804], font_size=8.8)

    add_heading(doc, "4. 个人总结", 1)
    add_body(
        doc,
        "本次学年设计的主要收获不是单独完成某个算法，而是建立了从数据获取、数据库维护、质量控制、模型分析、可视化到智能问数的完整工程意识。数据规模必须用 SQL 核验，采集失败必须有日志，模型结论必须同时给出样本量和误差指标，Agent 回答必须能追溯到工具结果。",
    )
    add_body(
        doc,
        "项目也暴露了需要继续改进的部分：不同来源的区县命名仍需进一步标准化，真实多期价格变化样本仍然有限，当前回归模型 R² 为 0.3847，只适合辅助解释挂牌单价，不能表述为精准预测成交价。后续工作应继续积累自然快照、增加分组交叉验证，并完成 Linux 服务器部署与恢复测试。",
    )

    add_heading(doc, "5. 后续工作计划（未完成项）", 1)
    add_bullets(
        doc,
        [
            "按真实采集周期继续积累 listing_snapshots，形成可用于趋势分析的多时间点数据。",
            "统一原始区县别名与标准化展示口径，减少分析任务中的重复分层标签。",
            "在 Linux 服务器上完成 Nginx、Gunicorn、systemd 和 MySQL 的真实部署验收。",
            "采用 GroupKFold 或时间切分复核回归模型，报告均值、标准差和数据版本。",
        ],
    )

    doc.core_properties.title = TITLE + " 学年设计日志"
    doc.core_properties.author = STUDENT
    doc.core_properties.subject = "学年设计Ⅱ过程记录"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / f"{TITLE}_学年设计日志.docx"
    doc.save(output)
    return output


def build_install() -> Path:
    doc = Document()
    configure_document(doc, TITLE + "｜安装与使用说明")
    add_cover(
        doc,
        "系统安装与使用说明",
        "Windows 本地运行与课程答辩操作手册",
        [
            ("系统名称", TITLE),
            ("适用版本", "2026年7月1日最终本地版"),
            ("运行方式", "Flask API + Vite/React + MySQL 8.x"),
            ("默认端口", "前端 5173；后端 5000；MySQL 3306"),
            ("数据口径", "二手房挂牌价/报价，不代表成交价"),
        ],
    )

    add_heading(doc, "1. 文档范围与系统边界", 1)
    add_body(
        doc,
        "本说明用于在 Windows 本地环境安装、启动和验收重庆二手房挂牌价分析系统。运行期数据库统一使用 MySQL 8.x；database 目录中的 SQL 用于恢复最终验收数据。旧 SQLite/CSV 只作为一次性冷启动来源，不作为运行期数据库。",
    )
    add_note(
        doc,
        "重要边界",
        "系统分析的是挂牌价/报价，不是成交价；模型用于解释和辅助估价，不承诺精准预测；采集遇到验证码或验证网关时记录失败，不进行绕过。",
    )

    add_heading(doc, "2. 交付目录", 1)
    add_table(
        doc,
        ["路径", "内容", "是否必需"],
        [
            ["Backend/", "Flask API、服务层、模型、爬虫、调度、分析和 Agent", "必需"],
            ["Frontend/", "Vite + React + ECharts 前端", "必需"],
            ["database/real_estate_final_20260701.sql", "最终 MySQL 数据与表结构", "首次恢复必需"],
            ["scripts/setup_local.ps1", "首次安装依赖并创建本地配置模板", "推荐"],
            ["scripts/start_local_dev.ps1", "同时启动后端和前端", "推荐"],
            [".env.example / Frontend/.env.example", "本地配置模板，不包含真实密钥", "必需"],
            ["docs/ / deliverables/", "测试证据和提交材料", "答辩使用"],
        ],
        [2900, 4584, 1700],
    )

    add_heading(doc, "3. 环境要求", 1)
    add_table(
        doc,
        ["组件", "最低/建议版本", "用途与检查命令"],
        [
            ["Windows", "Windows 10/11", "本地开发与 PowerShell 脚本"],
            ["Python", "3.8+", "python --version"],
            ["Node.js", "18+", "node --version"],
            ["npm", "随 Node.js 安装", "npm --version"],
            ["MySQL", "8.x", "mysql --version"],
            ["浏览器", "Chrome 或 Edge", "访问本地系统"],
            ["DeepSeek API Key", "可选", "启用智能问答、自然语言 SQL 和报告"],
            ["高德 JS API Key", "可选", "启用高德底图；无 Key 时仍可使用本地统计与其他页面"],
        ],
        [2100, 2100, 4984],
    )

    add_heading(doc, "4. 数据库初始化", 1)
    add_heading(doc, "4.1 创建数据库", 2)
    add_code(
        doc,
        'mysql -u root -p -e "CREATE DATABASE IF NOT EXISTS real_estate DEFAULT CHARACTER SET utf8mb4 COLLATE utf8mb4_unicode_ci;"',
        "在项目根目录执行",
    )
    add_heading(doc, "4.2 导入最终 SQL", 2)
    add_body(doc, "PowerShell 本身不直接支持传统的 < 输入重定向，建议通过 cmd /c 调用 MySQL 客户端：", first_indent=False)
    add_code(doc, 'cmd /c "mysql -u root -p real_estate < database\\real_estate_final_20260701.sql"')
    add_heading(doc, "4.3 导入后核验", 2)
    add_code(
        doc,
        """mysql -u root -p real_estate

SELECT COUNT(*) FROM listings WHERE status IN ('active','valid');
SELECT COUNT(DISTINCT district) FROM listings;
SELECT source, COUNT(*) FROM listings GROUP BY source ORDER BY COUNT(*) DESC;
SELECT COUNT(*) FROM listing_snapshots;""",
    )
    add_note(
        doc,
        "当前验收结果",
        "有效房源 100,897 条；Dashboard 标准化展示 38 个区县；fang 98,134 条、anjuke_mobile 2,753 条、lianjia 10 条；快照 100,914 条。",
        fill=PALE_GREEN,
    )

    add_heading(doc, "5. 本地配置", 1)
    add_heading(doc, "5.1 配置文件规则", 2)
    add_bullets(
        doc,
        [
            "根目录 .env：数据库、登录、采集、DeepSeek 和调度配置。",
            "Frontend/.env.local：高德 Web Key 与安全密钥等前端本地配置。",
            "data/cookies/：需要登录态的数据源 Cookie 文件。",
            "上述路径均被 .gitignore 排除，应保留在本机，禁止提交真实密钥和 Cookie。",
            "setup_local.ps1 只在文件不存在时从 example 创建，不会覆盖已有配置。",
        ],
    )
    add_heading(doc, "5.2 关键环境变量", 2)
    add_table(
        doc,
        ["变量", "作用", "说明"],
        [
            ["DATABASE_URL", "业务数据库连接", "必须指向 MySQL real_estate"],
            ["TEST_DATABASE_URL", "测试数据库连接", "建议使用独立 real_estate_test"],
            ["AUTH_ADMIN_USERNAME", "管理员账号", "默认 admin"],
            ["AUTH_ADMIN_PASSWORD", "管理员密码", "默认 swu@2026，正式环境应修改"],
            ["CRAWL_MAX_WORKERS", "采集最大并发", "建议 3—5，避免高频请求"],
            ["DEEPSEEK_ENABLED", "启用 Agent", "true 时还需有效 API Key"],
            ["DEEPSEEK_API_KEY", "DeepSeek 密钥", "只写入本地 .env"],
            ["SCHEDULER_ENABLED", "启用 APScheduler", "本地可按需开启"],
            ["INCREMENTAL_CRAWL_JOB_ENABLED", "启用定期增量采集", "需先确认采集源可用"],
            ["VITE_AMAP_WEB_KEY", "高德地图 Web Key", "写入 Frontend/.env.local"],
        ],
        [2700, 2900, 3584],
        font_size=8.8,
    )

    doc.add_page_break()
    add_heading(doc, "6. 安装与启动", 1)
    add_heading(doc, "6.1 推荐一键安装", 2)
    add_code(doc, "powershell -ExecutionPolicy Bypass -File .\\scripts\\setup_local.ps1")
    add_body(
        doc,
        "脚本会检查 Python 和 npm，在配置缺失时复制模板，然后安装 requirements.txt 与 Frontend/package-lock.json 中的依赖。已有 .env 和 Frontend/.env.local 会原样保留。",
    )
    add_heading(doc, "6.2 推荐同时启动前后端", 2)
    add_code(doc, "powershell -ExecutionPolicy Bypass -File .\\scripts\\start_local_dev.ps1")
    add_body(
        doc,
        "启动脚本会检查 5000/5173 端口，只把 Listen 状态视为占用；前端最多等待 20 秒。运行日志写入 .codex-run，服务成功后访问 http://127.0.0.1:5173。",
    )
    add_heading(doc, "6.3 手动启动", 2)
    add_code(
        doc,
        """# 终端 1：后端
python -m flask --app Backend.app run --debug --port 5000

# 终端 2：前端
cd Frontend
npm run dev -- --host 127.0.0.1 --port 5173""",
    )
    add_heading(doc, "6.4 启动检查", 2)
    add_code(
        doc,
        """curl http://127.0.0.1:5000/api/health
curl http://127.0.0.1:5000/api/overview

# 浏览器
http://127.0.0.1:5173""",
    )
    add_note(doc, "正常状态", "健康接口返回 status=healthy；前端首页显示 100,897 条房源和 38 个标准化区县。", fill=PALE_GREEN)

    add_heading(doc, "7. 登录与页面使用", 1)
    add_heading(doc, "7.1 登录", 2)
    add_table(
        doc,
        ["项目", "默认值", "说明"],
        [["地址", "http://127.0.0.1:5173", "本地前端"], ["用户名", "admin", "可在 .env 修改"], ["密码", "swu@2026", "正式使用前应修改"]],
        [2200, 2900, 4084],
    )
    add_heading(doc, "7.2 页面功能", 2)
    add_table(
        doc,
        ["页面", "主要操作", "数据来源/结果"],
        [
            ["首页总览", "查看 KPI、地图、区县排行、趋势、散点、总价与户型分布", "/api/overview 与 /api/charts/*"],
            ["房源数据管理", "关键词、区县、来源、价格和面积筛选；分页、详情、CSV 导出", "MySQL listings"],
            ["采集任务管理", "创建、运行、取消任务；查看进度、失败原因和页级日志", "crawl_tasks / crawl_logs"],
            ["数据清洗质量", "查看六维质量、异常样本、来源分层和清洗步骤", "data_quality_reports 与实时聚合"],
            ["分析建模", "运行 EDA、回归、参数搜索、聚类和异常检测；查看历史任务", "analysis_jobs / model_results"],
            ["智能问答与报告", "市场问数、模型解释、工具轨迹、自然语言 SQL、PDF 报告", "DeepSeek + ToolRegistry"],
            ["系统设置", "修改采集、调度和 DeepSeek 配置，执行连接测试", "system_settings 与本地环境变量"],
        ],
        [1900, 4300, 2984],
        font_size=8.8,
    )

    add_heading(doc, "8. 典型操作流程", 1)
    add_heading(doc, "8.1 创建采集任务", 2)
    add_numbered(
        doc,
        [
            "进入“采集任务管理”，点击新建任务。",
            "选择数据源和区县，首次测试建议每区 1 页、并发 1—3。",
            "创建后启动任务，观察 pending、running、success 或 partial_failed 状态。",
            "查看日志中的 URL、HTTP 状态、解析数量和失败原因。",
            "遇验证码或验证网关时停止重复尝试，不绕过验证。",
        ],
    )
    add_heading(doc, "8.2 执行分析任务", 2)
    add_numbered(
        doc,
        [
            "进入“分析建模”，选择重新训练或相应分析类型。",
            "等待任务完成后检查样本量、训练集、测试集和区县/来源覆盖。",
            "回归结果同时查看 MAE、RMSE、R² 和 MAPE，不只看单一最好指标。",
            "聚类用于市场画像；异常结果用于人工复核，不直接删除房源。",
        ],
    )
    add_heading(doc, "8.3 使用智能问答", 2)
    add_numbered(
        doc,
        [
            "在系统设置中启用 DeepSeek，填写 Key，并执行真实连接测试。",
            "进入“智能问答与报告”，提出区县统计、模型结果或采集状态问题。",
            "检查右侧执行轨迹，确认工具名、参数、结果和耗时。",
            "复杂聚合问题会调用 query_readonly_sql；生成 SQL 必须通过只读安全校验。",
            "若工具无数据，系统应说明数据不足，不应给出无法核验的具体数值。",
        ],
    )

    doc.add_page_break()
    add_heading(doc, "9. 验收与测试", 1)
    add_heading(doc, "9.1 自动验收命令", 2)
    add_code(
        doc,
        """# 后端语法和测试
python -m compileall Backend
python -m pytest Backend/tests -q

# 前端生产构建
cd Frontend
npm run build

# 服务启动后的核心接口 smoke
cd ..
python scripts/local_demo_smoke.py""",
    )
    add_heading(doc, "9.2 当前复验结果", 2)
    add_table(
        doc,
        ["验收项", "结果", "说明"],
        [
            ["后端测试", "82 passed", "覆盖认证、房源、图表、采集、质量、分析、Agent、调度和设置"],
            ["前端构建", "通过", "Vite 生产构建成功；存在主包较大的非阻塞提示"],
            ["核心接口 smoke", "通过", "health、overview、listings、quality、analysis、settings、agent/tools"],
            ["DeepSeek 连接", "通过", "真实轻量请求成功"],
            ["自然语言 SQL", "通过", "真实调用 query_readonly_sql，返回受控聚合结果"],
            ["PC 页面", "通过", "1920×900 页面渲染，控制台无相关错误"],
        ],
        [2200, 1600, 5384],
        font_size=8.8,
    )

    add_heading(doc, "10. 常见问题", 1)
    add_table(
        doc,
        ["现象", "排查与处理"],
        [
            ["数据库连接失败", "确认 MySQL 服务、3306 端口、real_estate 库和 DATABASE_URL；检查账号权限。"],
            ["导入 SQL 失败", "确认 mysql 命令可用；在 PowerShell 中使用 cmd /c 重定向命令；检查磁盘空间。"],
            ["5000 或 5173 被占用", "使用 Get-NetTCPConnection -State Listen 查看占用进程，确认后再停止；不要直接结束不明进程。"],
            ["前端显示接口失败", "先访问 /api/health；检查 Vite 代理、后端日志和浏览器控制台。"],
            ["采集任务 partial_failed", "查看 crawl_logs；网站结构变化、限流或验证网关均可能导致部分失败。"],
            ["Agent 提示不可用", "检查 DEEPSEEK_ENABLED、API Key、模型名和网络；在设置页运行连接测试。"],
            ["地图不显示高德底图", "检查 Frontend/.env.local 的 Web Key 和安全码；其他统计与业务页面仍可运行。"],
            ["分析结果为空", "确认存在质量分不低于 80 且价格、面积有效的数据，再创建分析任务。"],
            ["启动脚本返回失败", "查看 .codex-run 中后端和前端 err.log；首次启动 Vite 最多等待 20 秒。"],
        ],
        [2800, 6384],
        font_size=8.8,
    )

    add_heading(doc, "11. 安全、数据与发布边界", 1)
    add_bullets(
        doc,
        [
            "不得提交 .env、Frontend/.env.local、data/cookies、API Key、数据库密码和真实 Token。",
            "Agent 不直接持有数据库连接，不执行用户原样输入的 SQL；只读 SQL 工具限制业务表、单条 SELECT、100 行、5 秒超时和只读事务。",
            "采集任务限制并发和请求间隔，遇验证码停止，不做强对抗反爬。",
            "所有房价均为挂牌价/报价；模型重要性表示统计关联或模型贡献，不代表因果关系。",
            "当前文档只证明 Windows 本地运行；Nginx、Gunicorn、systemd 和阿里云公网部署属于后续方案，不能表述为已上线。",
        ],
    )

    doc.core_properties.title = TITLE + " 安装与使用说明"
    doc.core_properties.author = STUDENT
    doc.core_properties.subject = "学年设计Ⅱ系统安装与使用电子文档"
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    output = OUTPUT_DIR / f"{TITLE}_安装与使用说明.docx"
    doc.save(output)
    return output


def main():
    log_path = build_log()
    install_path = build_install()
    print(log_path)
    print(install_path)


if __name__ == "__main__":
    main()
